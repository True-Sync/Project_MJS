# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "Project_MJS_Cinematic_Architecture.docx"

FONT = "Malgun Gothic"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(24, 32, 43)
MUTED = RGBColor(90, 96, 106)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "D9E0EA"


def set_east_asia_font(run, name: str = FONT) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_in):
                cell.width = Inches(widths_in[idx])
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.tcW
                if tc_w is not None:
                    tc_w.set(qn("w:type"), "dxa")
                    tc_w.set(qn("w:w"), str(int(widths_in[idx] * 1440)))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_border_bottom(paragraph, color: str = "B8C7D9", size: str = "8") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def style_run(run, size=11, color=INK, bold=False, italic=False) -> None:
    set_east_asia_font(run)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_para(doc, text: str, size=11, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    style_run(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph()
    if level == 1:
        size, color, before, after = 16, BLUE, 18, 10
    elif level == 2:
        size, color, before, after = 13, BLUE, 14, 7
    else:
        size, color, before, after = 12, DARK_BLUE, 10, 5
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=size, color=color, bold=True)
    return p


def add_bullet(doc, text: str, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    style_run(run)
    return p


def add_numbered(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    style_run(run)
    return p


def add_callout(doc, title: str, body: str, fill: str = LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run(title)
    style_run(r_title, size=11, color=DARK_BLUE, bold=True)
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_after = Pt(0)
    p_body.paragraph_format.line_spacing = 1.25
    r_body = p_body.add_run(body)
    style_run(r_body, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_key_value_table(doc, rows: list[tuple[str, str]], widths=(1.45, 5.05)):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, list(widths))
    hdr = table.rows[0].cells
    hdr[0].text = "항목"
    hdr[1].text = "설명"
    for cell in hdr:
        set_cell_shading(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            for r in p.runs:
                style_run(r, size=10.5, color=DARK_BLUE, bold=True)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for idx, cell in enumerate(cells):
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.2
                for r in p.runs:
                    style_run(r, size=10.2, color=DARK_BLUE if idx == 0 else INK, bold=(idx == 0))
            set_cell_margins(cell, top=90, bottom=90)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_matrix_table(doc, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                style_run(r, size=10, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.18
                for r in p.runs:
                    style_run(r, size=9.8, color=INK)
            set_cell_margins(cells[i], top=85, bottom=85)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_code_block(doc, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FC")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for line_idx, line in enumerate(text.splitlines()):
        if line_idx:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.size = Pt(9.2)
        r.font.color.rgb = RGBColor(45, 55, 72)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def set_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    for style_name in ("Normal", "List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(11)
        style.font.color.rgb = INK
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "Project_MJS Cinematic Architecture"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs:
        style_run(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Generated architecture note"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in footer.runs:
        style_run(r, size=9, color=MUTED)


def build_doc() -> None:
    doc = Document()
    set_doc_styles(doc)

    add_para(doc, "Project_MJS", size=10.5, color=MUTED, bold=True, after=2)
    title = add_para(doc, "시네마틱 시스템 구조 설명서", size=24, color=RGBColor(0, 0, 0), bold=True, after=4)
    subtitle = add_para(
        doc,
        "스킬 컷신, 일반 컷신, 레벨 트리거 연출을 하나의 재생 구조로 다루기 위한 아키텍처 노트",
        size=12.5,
        color=MUTED,
        after=14,
    )
    set_paragraph_border_bottom(subtitle)
    add_para(doc, "인코딩: UTF-8 / 문서 유형: compact_reference_guide / 폰트 오버라이드: Malgun Gothic", size=9.5, color=MUTED, italic=True, after=14)

    add_callout(
        doc,
        "핵심 요약",
        "현재 구조의 목표는 스킬, 궁극기, 일반 컷신, 레벨 트리거가 모두 같은 DirectorSubsystem에 재생 요청을 보내고, 입력 잠금이나 월드 정지는 각 액터의 ParticipantComponent가 독립적으로 처리하게 만드는 것이다. 즉 요청자와 반응자를 분리해서 의존성을 낮춘다.",
        LIGHT_BLUE,
    )

    add_heading(doc, "1. 설계 목표", 1)
    for item in [
        "스킬 시퀀스와 일반 컷신 시퀀스를 같은 재생 파이프라인으로 처리한다.",
        "플레이어 컨트롤러, 캐릭터 이동, 적 AI, 애니메이션 정지 같은 정책을 시퀀스 재생 코드에서 분리한다.",
        "명조식 궁극기 컷인처럼 다른 화면/카메라 연출이 나오는 동안 참가 액터를 선택적으로 멈출 수 있게 한다.",
        "트리거 액터, 스킬 컴포넌트, 대화 이벤트가 서로를 몰라도 FCinematicPlaybackRequest만으로 시네마틱을 발동할 수 있게 한다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. 전체 구조", 1)
    add_code_block(
        doc,
        "스킬/궁극기/트리거/대화\n"
        "        |\n"
        "        v\n"
        "FCinematicPlaybackRequest\n"
        "        |\n"
        "        v\n"
        "UCinematicDirectorSubsystem\n"
        "   |                 |\n"
        "   v                 v\n"
        "LevelSequencePlayer  ICinematicParticipant 알림\n"
        "                     |\n"
        "                     v\n"
        "        UCinematicParticipantComponent\n"
        "        입력/이동/Tick/애니메이션 정지 및 복구",
    )
    add_para(
        doc,
        "구조의 중심은 UCinematicDirectorSubsystem이다. 요청자는 Director에게 재생 요청만 보내고, Director는 Level Sequence 재생과 참가자 시작/종료 알림을 담당한다. 실제로 플레이어 입력을 막거나 적을 정지시키는 처리는 참가자 컴포넌트가 각자 수행한다.",
    )

    add_heading(doc, "3. 주요 클래스 책임", 1)
    add_matrix_table(
        doc,
        ["클래스/구조체", "책임", "의존성 방향"],
        [
            ["FCinematicPlaybackRequest", "외부 요청 데이터. 어떤 Sequence를 누가, 누구 중심으로, 어떤 범위에 적용할지 담는다.", "요청자 -> Director"],
            ["FCinematicPlaybackContext", "재생 중 참가자에게 전달되는 런타임 정보. SequenceActor, PlayerController, SubjectActor 등을 포함한다.", "Director -> Participant"],
            ["UCinematicDirectorSubsystem", "월드 단위 재생 관리자. SequencePlayer 생성, 참가자 수집, 시작/종료 알림, ViewTarget 복구를 처리한다.", "중앙 조정자"],
            ["ICinematicParticipant", "시네마틱 시작/종료 이벤트를 받는 인터페이스. 액터나 컴포넌트가 선택적으로 구현한다.", "Director -> 구현체"],
            ["UCinematicParticipantComponent", "입력 잠금, 이동 정지, Tick 정지, 애니메이션 정지를 옵션으로 수행한다.", "구현체 내부 정책"],
            ["UCinematicActionComponent", "액터 내부에서 간단히 시네마틱을 요청하는 얇은 래퍼. 직접 잠금 처리를 하지 않는다.", "액터 -> Director"],
            ["ACinematicTriggerActor", "레벨 오버랩 조건으로 요청을 만든다. 대상 액터에 ActionComponent가 없어도 동작한다.", "트리거 -> Director"],
        ],
        [1.65, 3.55, 1.3],
    )

    add_heading(doc, "4. 요청 데이터와 컨텍스트", 1)
    add_key_value_table(
        doc,
        [
            ["Sequence", "실제로 재생할 ULevelSequence 에셋. null이면 요청은 거부된다."],
            ["InstigatorActor", "시네마틱을 발동한 액터. 예: 플레이어, 스킬 사용자, 트리거를 밟은 액터."],
            ["SubjectActor", "시네마틱의 중심 대상. 카메라나 연출의 주인공이 되는 액터."],
            ["PlayerController", "입력 잠금과 ViewTarget 복구에 사용할 컨트롤러. 비워두면 Director가 자동 탐색한다."],
            ["AdditionalParticipants", "Instigator/Subject 외에 직접 참가자로 포함할 액터 목록."],
            ["bAffectAllParticipants", "true면 월드의 모든 ICinematicParticipant 구현체를 수집한다. false면 명시된 대상만 반응한다."],
            ["bRestoreViewTarget", "시네마틱 종료 시 이전 ViewTarget으로 돌아갈지 결정한다."],
            ["bStopPreviousCinematic", "기존 시네마틱이 재생 중일 때 새 요청으로 교체할지 결정한다."],
            ["BlendOutTime", "기존 ViewTarget 복구 시 사용할 카메라 블렌드 시간."],
        ],
    )

    add_heading(doc, "5. 재생 흐름", 1)
    for step in [
        "스킬 컴포넌트, 트리거 액터, 대화 이벤트 등이 FCinematicPlaybackRequest를 만든다.",
        "UCinematicDirectorSubsystem::PlayCinematic()이 요청을 검증하고 기존 재생 중인 시네마틱을 처리한다.",
        "Director가 ULevelSequencePlayer와 런타임 ALevelSequenceActor를 생성한다.",
        "Director가 PlayerController, 이전 ViewTarget, 참가자 목록을 수집한다.",
        "참가자들에게 OnCinematicStarted(Context)를 호출한다.",
        "UCinematicParticipantComponent가 옵션에 따라 입력, 이동, Tick, 애니메이션을 정지한다.",
        "Level Sequence가 재생된다.",
        "시퀀스 종료 이벤트가 오면 Director가 참가자들에게 OnCinematicEnded(Context)를 호출한다.",
        "참가자들이 자신이 잠근 상태를 복구하고, Director가 ViewTarget과 SequenceActor를 정리한다.",
    ]:
        add_numbered(doc, step)

    add_heading(doc, "6. 스킬/궁극기 시나리오", 1)
    add_para(
        doc,
        "스킬 시스템이 추가되면 SkillComponent는 전투 상태, 쿨타임, 게이지, 캔슬 가능 여부만 판단한다. 연출이 필요한 궁극기라면 몽타주 재생과 함께 FCinematicPlaybackRequest를 Director에 전달한다. 데미지 판정은 기존 원칙대로 AnimNotify 또는 HitWindow 쪽에 남기고, Level Sequence는 카메라 컷, FOV, 포스트 프로세스, 연출용 VFX 같은 시각 레이어를 담당한다.",
    )
    add_code_block(
        doc,
        "SkillInput\n"
        "  -> SkillComponent: 쿨타임/게이지/캔슬 가능 여부 검사\n"
        "  -> AnimMontage 재생\n"
        "  -> FCinematicPlaybackRequest 생성\n"
        "  -> CinematicDirectorSubsystem 재생 요청\n"
        "  -> ParticipantComponent들이 입력/이동/적 정지 처리\n"
        "  -> Notify에서 실제 타격/사운드/VFX 처리",
    )

    add_heading(doc, "7. 일반 컷신/레벨 트리거 시나리오", 1)
    add_para(
        doc,
        "ACinematicTriggerActor는 레벨에 배치해서 사용할 수 있는 기본 트리거다. 박스 오버랩, 1회성 실행, RequiredActorClass 필터, 활성화 여부를 가진다. 중요한 점은 대상 액터가 UCinematicActionComponent를 반드시 가질 필요가 없다는 것이다. 트리거는 DirectorSubsystem에 직접 요청을 보내므로 일반 컷신, 보스 등장, 지역 진입 연출, 대화 전환 등에 같은 구조를 사용할 수 있다.",
    )

    add_heading(doc, "8. 참가자 정지 정책", 1)
    add_matrix_table(
        doc,
        ["옵션", "처리 내용", "권장 사용"],
        [
            ["bLockPlayerMoveInput", "PlayerController의 이동 입력을 무시하도록 설정한다.", "플레이어 조작을 막는 컷신"],
            ["bLockPlayerLookInput", "PlayerController의 시점 입력을 무시하도록 설정한다.", "카메라 컷이 중요한 연출"],
            ["bDisableCharacterMovement", "CharacterMovement를 정지하고 종료 시 이전 MovementMode로 복구한다.", "플레이어/적 위치 고정"],
            ["bDisableOwnerTick", "소유 액터 Tick을 끈다.", "비전투 NPC, 단순 장치 정지"],
            ["bDisableComponentTicks", "소유 액터의 다른 컴포넌트 Tick을 끈다.", "AI/로직 컴포넌트가 Tick 기반일 때"],
            ["bPauseSkeletalAnimations", "SkeletalMeshComponent의 애니메이션을 일시정지한다.", "배경 캐릭터를 완전히 얼리고 싶을 때"],
        ],
        [1.75, 3.15, 1.6],
    )

    add_callout(
        doc,
        "중요한 설계 포인트",
        "월드 전체 Pause나 GlobalTimeDilation 0은 시퀀서, Niagara, 오디오와 충돌할 수 있다. 현재 구조는 전체 월드를 강제로 멈추기보다, 멈춰야 할 액터에 ParticipantComponent를 붙여 선택적으로 반응하게 하는 쪽을 기본 전략으로 삼는다.",
        LIGHT_BLUE,
    )

    add_heading(doc, "9. 현재 플레이어 구성", 1)
    add_key_value_table(
        doc,
        [
            ["AttackComponent", "기본 공격과 콤보 몽타주를 처리한다. 시네마틱 구조와 직접 결합하지 않는다."],
            ["DodgeComponent", "회피 몽타주를 처리한다. 시네마틱 시작/종료 정책과 분리되어 있다."],
            ["CinematicActionComponent", "플레이어 내부에서 시네마틱 요청을 보낼 때 사용하는 래퍼다."],
            ["CinematicParticipantComponent", "시네마틱 중 플레이어 입력과 이동 정지를 담당한다."],
        ],
    )

    add_heading(doc, "10. 의존성 감소 효과", 1)
    for item in [
        "트리거 액터는 플레이어 클래스나 스킬 클래스를 모른다.",
        "스킬 컴포넌트는 적 AI나 카메라 복구 로직을 몰라도 된다.",
        "참가자 컴포넌트는 요청자가 스킬인지 레벨 트리거인지 몰라도 된다.",
        "DirectorSubsystem만 시퀀스 재생과 참가자 통지를 중앙에서 관리한다.",
        "새로운 적, NPC, 보스는 ParticipantComponent를 붙이는 것만으로 컷신 반응 대상이 될 수 있다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11. 남은 개선점", 1)
    add_matrix_table(
        doc,
        ["개선 항목", "필요한 이유", "추천 방향"],
        [
            ["입력 잠금 토큰화", "현재는 SetIgnoreMoveInput(false) 복구가 다른 시스템의 잠금을 풀 위험이 있다.", "InputLockSubsystem 또는 InputLockComponent를 만들어 핸들 기반 Acquire/Release 구조로 전환"],
            ["AI 정지 정책 분리", "Enemy/AIController 구현이 생기면 단순 Tick 정지만으로 부족할 수 있다.", "AIController/BrainComponent 정지를 처리하는 별도 Participant 옵션 추가"],
            ["시퀀스 바인딩", "LevelSequence 내부 카메라/캐릭터 바인딩을 런타임 액터와 연결해야 할 수 있다.", "Request에 BindingOverrides 또는 SubjectBindingName 추가"],
            ["네트워크 고려", "멀티플레이나 리플레이를 고려하면 로컬 컷신과 서버 권한 판정이 분리되어야 한다.", "Director 요청을 로컬/서버 권한으로 구분"],
            ["스킬 시스템 연결", "궁극기 입력, 쿨타임, 캔슬 윈도우와 아직 연결되지 않았다.", "SkillComponent에서 Request를 만들고 Montage/Notify와 병행"],
        ],
        [1.45, 2.55, 2.5],
    )

    add_heading(doc, "12. 사용 예시", 1)
    add_para(doc, "레벨 컷신은 ACinematicTriggerActor를 배치하고 Sequence, RequiredActorClass, bTriggerOnce를 설정한다. 플레이어와 멈춰야 할 적/NPC에는 UCinematicParticipantComponent를 붙인다.")
    add_para(doc, "스킬 컷신은 SkillComponent에서 조건 검사를 끝낸 뒤 FCinematicPlaybackRequest를 생성해 DirectorSubsystem에 넘긴다. 이때 bAffectAllParticipants를 true로 두면 참가자 컴포넌트가 붙은 모든 액터가 반응하고, false로 두면 요청에 명시된 대상만 반응한다.")

    add_heading(doc, "13. 파일 위치", 1)
    add_key_value_table(
        doc,
        [
            ["Public/Cinematic/CinematicTypes.h", "요청/컨텍스트 데이터 정의"],
            ["Public/Cinematic/CinematicDirectorSubsystem.h", "월드 단위 시네마틱 관리자 선언"],
            ["Private/Cinematic/CinematicDirectorSubsystem.cpp", "시퀀스 재생, 참가자 수집, 종료 복구 구현"],
            ["Public/Cinematic/CinematicParticipant.h", "참가자 인터페이스"],
            ["Public/Cinematic/CinematicParticipantComponent.h", "참가자 컴포넌트 선언 및 옵션"],
            ["Private/Cinematic/CinematicParticipantComponent.cpp", "입력/이동/Tick/애니메이션 정지 구현"],
            ["Public/Private/Cinematic/CinematicActionComponent", "액터 내부 요청용 래퍼"],
            ["Public/Private/Cinematic/CinematicTriggerActor", "레벨 오버랩 트리거"],
        ],
    )

    add_callout(
        doc,
        "추천 결론",
        "현재 구조는 Project_MJS에서 스킬 컷신과 일반 컷신을 함께 수용하기 위한 기반으로 적절하다. 다음 단계는 InputLockSubsystem을 추가해서 입력 잠금 복구를 더 안전하게 만들고, 이후 SkillComponent가 DirectorSubsystem에 요청을 보내도록 연결하는 것이다.",
        LIGHT_BLUE,
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
